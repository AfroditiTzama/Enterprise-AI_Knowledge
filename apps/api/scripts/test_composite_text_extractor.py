import asyncio
from io import BytesIO
from pathlib import Path

from docx import Document as OpenXmlDocument

from knowledge_assistant.infrastructure.documents.text_extractors.composite_text_extractor import (
    CompositeDocumentTextExtractor,
)


PDF_PATH = Path(
    "storage/documents/de54408d24c9433694d59aacae1f60bc.pdf"
)


def create_test_docx() -> bytes:
    document = OpenXmlDocument()

    document.add_heading(
        "Enterprise AI Knowledge Assistant",
        level=1,
    )

    document.add_paragraph(
        "This DOCX file was processed by the composite extractor."
    )

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()


async def test_pdf(
    extractor: CompositeDocumentTextExtractor,
) -> None:
    if not PDF_PATH.exists():
        raise FileNotFoundError(
            f"Test PDF was not found: {PDF_PATH}"
        )

    extracted_document = await extractor.extract(
        file_content=PDF_PATH.read_bytes(),
        original_filename="5_AI_PROVIDERS.pdf",
        content_type="application/pdf",
    )

    print("PDF extraction successful")
    print(
        f"Segments: {len(extracted_document.segments)}"
    )
    print(
        f"Characters: {len(extracted_document.full_text)}"
    )


async def test_docx(
    extractor: CompositeDocumentTextExtractor,
) -> None:
    extracted_document = await extractor.extract(
        file_content=create_test_docx(),
        original_filename="test_document.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    print("\nDOCX extraction successful")
    print(
        f"Segments: {len(extracted_document.segments)}"
    )
    print(extracted_document.full_text)


async def test_txt(
    extractor: CompositeDocumentTextExtractor,
) -> None:
    text = (
        "Enterprise AI Knowledge Assistant\n\n"
        "Αυτό είναι ένα δοκιμαστικό αρχείο TXT."
    )

    extracted_document = await extractor.extract(
        file_content=text.encode("utf-8"),
        original_filename="test_document.txt",
        content_type="text/plain",
    )

    print("\nTXT extraction successful")
    print(
        f"Segments: {len(extracted_document.segments)}"
    )
    print(extracted_document.full_text)


async def test_unsupported_file(
    extractor: CompositeDocumentTextExtractor,
) -> None:
    try:
        await extractor.extract(
            file_content=b"fake image content",
            original_filename="image.png",
            content_type="image/png",
        )
    except ValueError as exc:
        print("\nUnsupported file test successful")
        print(exc)
        return

    raise AssertionError(
        "Unsupported file type was not rejected."
    )


async def main() -> None:
    extractor = CompositeDocumentTextExtractor()

    await test_pdf(extractor)
    await test_docx(extractor)
    await test_txt(extractor)
    await test_unsupported_file(extractor)


if __name__ == "__main__":
    asyncio.run(main())