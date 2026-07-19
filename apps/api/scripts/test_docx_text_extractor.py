import asyncio
from io import BytesIO

from docx import Document as OpenXmlDocument

from knowledge_assistant.infrastructure.documents.text_extractors.docx_text_extractor import (
    DOCXTextExtractor,
)


def create_test_docx() -> bytes:
    document = OpenXmlDocument()

    document.add_heading(
        "Enterprise AI Knowledge Assistant",
        level=1,
    )

    document.add_paragraph(
        "This document is used to test DOCX text extraction."
    )

    table = document.add_table(rows=2, cols=3)

    table.cell(0, 0).text = "Provider"
    table.cell(0, 1).text = "Region"
    table.cell(0, 2).text = "SLA"

    table.cell(1, 0).text = "ExampleAI"
    table.cell(1, 1).text = "Europe"
    table.cell(1, 2).text = "Available"

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()


async def main() -> None:
    file_content = create_test_docx()

    extractor = DOCXTextExtractor()

    extracted_document = await extractor.extract(
        file_content=file_content,
        original_filename="test_document.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    print(
        f"Extracted segments: "
        f"{len(extracted_document.segments)}"
    )

    print(  
        f"Total characters: "
        f"{len(extracted_document.full_text)}"
    )

    print("\nExtracted content:")
    print("-" * 80)

    for index, segment in enumerate(
        extracted_document.segments,
        start=1,
    ):
        print(f"\nSegment {index}")
        print(f"Page number: {segment.page_number}")
        print(segment.text)

    print("-" * 80)


if __name__ == "__main__":
    asyncio.run(main())